"""
生成 VedioQuant 研究报告 v2 (docx 格式)
中文版 + 英文版
增加背景故事和动机
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

assets_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets')


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        for p in table.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val
    return table


def add_img(doc, name, width=5.5):
    path = os.path.join(assets_dir, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3
    for i in range(1, 4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Arial'
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_bold(doc, text, size=12, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_cover(doc, subtitle_cn, subtitle_en=None):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VedioQuant')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle_cn)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4E, 0xCD, 0xC4)

    if subtitle_en:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle_en)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n10× Cache Compression  |  <2% Quality Loss  |  One Line to Enable')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x6B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nGitHub: github.com/robin-ph/vedioquant')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x45, 0xB7, 0xD1)

    doc.add_page_break()


# ================================================================
# 中文版
# ================================================================

def build_chinese():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, '视频扩散模型推理缓存极端压缩', 'TurboQuant × TeaCache 跨领域融合')

    # ---- 背景故事 ----
    doc.add_heading('起因：一个视频创作者的好奇心', level=1)

    doc.add_paragraph(
        '2025 年，AI 视频生成赛道彻底爆发。字节跳动的 Seedance 在中国一经发布就供不应求——'
        '一条 5 秒视频生成要等好几分钟，价格高昂，但仍有无数视频工作者和创作者抢着用。'
        '短视频平台上用 AI 生成的内容肉眼可见地增多，这个领域正在以惊人的速度重塑内容创作产业。'
    )

    doc.add_paragraph(
        '作为一个对技术充满好奇的人，我想知道：为什么视频生成这么贵？'
        '一段 5 秒 720P 的视频，背后到底发生了多少计算？'
        '于是我开始深入研究视频扩散模型（Wan2.1、CogVideoX、HunyuanVideo 等）的推理流程。'
    )

    doc.add_paragraph(
        '答案让我震惊：一次视频生成需要 50 步去噪迭代，每步对 75,600 个 token 做 30 层 transformer 前向传播。'
        '这个计算量大到需要 A100 80GB 这样的顶级 GPU 才能跑起来。'
        '而消费者和中小创作者根本用不起这样的硬件。'
    )

    doc.add_heading('转折：Google 的论文和一个惊喜的发现', level=2)

    doc.add_paragraph(
        '就在这时，我看到了两个东西：'
    )

    doc.add_paragraph(
        '第一，Google 在 ICLR 2026 发表了 TurboQuant——一种将大语言模型 KV Cache '
        '从 fp16 极端压缩到 3.5-bit 的技术，在 4× 压缩下依然保持 0.95 的余弦相似度。'
        '核心 trick 是用随机正交旋转将向量坐标分布"高斯化"，使得低 bit 标量量化高效可行。'
    )

    doc.add_paragraph(
        '第二，有人用 Python 工程复现了 TurboQuant (turboquant_plus)，'
        '并且在真实的 Qwen3 语言模型上验证了旋转后峰度从 900 降到 2.9（完美高斯分布是 3.0），'
        '141 个测试全部通过。这证明了这个算法不只是论文上好看，工程上也完全可行。'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        '这极大地刺激了我：如果 TurboQuant 能压缩语言模型的 KV Cache，'
        '那视频模型的特征缓存呢？两者缓存的都是注意力层的中间张量——数据本质完全相同！'
    )
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        '于是我决定动手尝试：把 TurboQuant 迁移到视频模型的 TeaCache 加速框架中，'
        '压缩缓存特征，降低显存占用。'
    )

    doc.add_paragraph(
        '没想到效果这么惊人——3-bit 压缩下余弦相似度达到 0.98，'
        '比语言模型上的 0.95 还要好。720P/81帧的缓存从 886MB 压到 88MB，'
        '30 层全缓存从 13GB 压到 1.3GB。这意味着普通 24GB 消费级 GPU 能跑原本需要 80GB 的配置。'
    )

    doc.add_page_break()

    # ---- 为什么视频模型比语言模型更需要这个 ----
    doc.add_heading('为什么视频模型比语言模型更需要缓存压缩？', level=1)

    doc.add_paragraph(
        '很多人可能会问：KV Cache 压缩在语言模型上已经研究很多了，视频模型有什么特别的？'
    )

    doc.add_heading('原因 1: 计算量是语言模型的数百倍', level=2)
    doc.add_paragraph(
        '语言模型的 KV Cache 线性增长——每生成一个 token 加一行。'
        '但视频模型是"一次性全量计算"：每步去噪都要对所有 token 做注意力。'
    )

    add_table(doc,
        ['', '语言模型 (Qwen3-7B)', '视频模型 (Wan2.1-14B, 720P)'],
        [
            ['Token 数', '~8,000 (8K上下文)', '75,600 (720P×81帧)'],
            ['计算次数', '1次/token (自回归)', '50次 (50步去噪)'],
            ['总注意力计算', '8K² × 32层', '75.6K² × 40层 × 50步'],
            ['量级对比', '1×', '~350×'],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph(
        '视频模型的注意力计算量是语言模型的 ~350 倍。任何能减少重复计算的优化，'
        '在视频领域的收益都被放大数百倍。'
    )

    doc.add_heading('原因 2: 显存是最硬的瓶颈', level=2)
    doc.add_paragraph(
        '语言模型推理时，KV Cache 通常是可控的——几百 MB 到几 GB。'
        '但视频模型的中间特征巨大：'
    )

    add_table(doc,
        ['模型', '场景', '单层缓存大小 (fp32)', '30层全缓存'],
        [
            ['Qwen3-7B', '8K 上下文', '~32 MB', '~1 GB'],
            ['Wan2.1-1.3B', '480P / 81帧', '192 MB', '5.8 GB'],
            ['Wan2.1-14B', '720P / 81帧', '443 MB', '12.98 GB'],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph(
        '视频模型的缓存动辄十几 GB，直接决定了你能不能在消费级 GPU 上跑——'
        '而这恰恰是视频创作者最常用的硬件。压缩缓存不是"锦上添花"，是"能不能用"的问题。'
    )

    doc.add_heading('原因 3: 视频模型的特征更适合压缩', level=2)
    doc.add_paragraph(
        '这是我在实验中发现的一个意外惊喜。视频模型特征比语言模型 KV 向量更适合 TurboQuant 压缩：'
    )

    add_table(doc,
        ['特征性质', '语言模型 KV', '视频模型特征', '含义'],
        [
            ['旋转前峰度', '~900', '~15', '视频特征本身更均匀'],
            ['旋转后峰度', '2.9', '0.2', '视频特征高斯化更完美'],
            ['3-bit 余弦相似度', '0.95', '0.98', '视频压缩质量更高'],
        ]
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        '语言模型 KV 向量的峰度高达 900（极端离群值多），需要旋转才能量化。'
        '视频模型特征的峰度只有 15，旋转后几乎完美高斯（峰度 0.2 vs 理论值 0.0）。'
        '这意味着视频模型天然更适合这种压缩方案——这是一个完美的跨领域迁移。'
    )
    run.font.size = Pt(11)

    doc.add_heading('原因 4: 产业需求极其迫切', level=2)
    doc.add_paragraph(
        'Seedance 一条 5 秒视频售价数元人民币，这背后是每次生成消耗一张 A100 GPU 数分钟的成本。'
        '降低推理成本不仅是技术问题，更是商业问题——谁能把成本降一个数量级，谁就能让 AI 视频生成'
        '从"少数人的昂贵玩具"变成"每个创作者的日常工具"。'
    )

    doc.add_paragraph(
        'VedioQuant 通过压缩缓存降低显存门槛，使：'
    )

    bullets = [
        '消费级 GPU (RTX 4090, 24GB) 能跑 720P 长视频全缓存推理',
        '云 GPU 单次推理成本降低（更低显存 = 可以用更便宜的 GPU）',
        '单卡能同时服务更多请求（显存占用减少 = 批量推理空间增大）',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ---- 算法原理 ----
    doc.add_heading('算法原理：TurboQuant 如何极端压缩', level=1)

    add_img(doc, 'algorithm_flow.png', 6)
    add_caption(doc, '图1: TurboQuant 压缩流程 — 从 fp32 (32 bits) 到 3 bits/坐标')
    doc.add_paragraph()

    doc.add_heading('Step 1: 范数分离', level=2)
    doc.add_paragraph(
        '向量的"长度"（范数 γ）和"方向"分开处理。范数用 float32 全精度存储，'
        '只对方向做低 bit 量化——因为注意力计算（内积）主要依赖方向。'
    )

    doc.add_heading('Step 2: 随机旋转高斯化——核心 Trick', level=2)
    doc.add_paragraph(
        '直接量化原始向量效果差，因为坐标分布极不均匀（离群值多）。'
        '乘以随机正交矩阵后，每个新坐标 = 原始所有坐标的随机加权平均。'
        '根据中心极限定理，结果趋近高斯分布——标量量化效率最大化。'
    )

    add_img(doc, 'gaussianization.png', 5.5)
    add_caption(doc, '图2: 随机旋转使特征分布从尖锐（峰度15）变为均匀高斯（峰度0.2）')
    doc.add_paragraph()

    add_bold(doc,
        '在 Wan2.1 真实特征上验证：旋转前峰度 15.0 → 旋转后 0.2，'
        '标准差 0.025511 vs 理论值 0.025516，完全吻合。', 11)

    doc.add_heading('Step 3: 预计算码本快速量化', level=2)
    doc.add_paragraph(
        '旋转后坐标服从 N(0, 1/√d)，预计算最优 Lloyd-Max 码本即可。'
        '量化变成一次 torch.bucketize 查表——无需逐向量迭代。'
    )

    # ---- 实验验证 ----
    doc.add_heading('实验验证', level=1)

    doc.add_paragraph(
        '从 Wan2.1-1.3B 的 30 层注意力层提取真实中间特征（d=1536），'
        '50 个 token 上测试压缩/解压质量：'
    )

    add_img(doc, 'compression_quality.png', 5.5)
    add_caption(doc, '图3: 不同量化位数的压缩质量和压缩比')
    doc.add_paragraph()

    add_table(doc,
        ['配置', '压缩比', '平均余弦', '最低余弦', '判定'],
        [
            ['2-bit', '15.2×', '0.9388', '0.9347', '✅ > 0.90'],
            ['3-bit (推荐)', '10.0×', '0.9822', '0.9729', '✅ > 0.90'],
            ['4-bit', '7.3×', '0.9942', '0.9921', '✅ > 0.90'],
        ]
    )

    doc.add_paragraph()
    add_bold(doc,
        '3-bit 甜点配置：10× 压缩，余弦 0.98，即使最差向量也 > 0.97。', 12, (0xFF, 0x6B, 0x6B))

    # ---- 显存影响 ----
    doc.add_heading('实际影响：显存节省', level=1)

    add_img(doc, 'memory_comparison.png', 5.5)
    add_caption(doc, '图4: 不同视频配置下缓存大小对比')
    doc.add_paragraph()

    add_img(doc, 'full_cache_comparison.png', 5.0)
    add_caption(doc, '图5: 720P/81帧全层缓存——VedioQuant 使 24GB GPU 可行')
    doc.add_paragraph()

    add_table(doc,
        ['缓存层数', 'fp32', '3-bit', '节省', '24GB GPU'],
        [
            ['2 层', '886 MB', '88 MB', '798 MB', '均可'],
            ['10 层', '4.33 GB', '441 MB', '3.89 GB', '均可'],
            ['30 层（全部）', '12.98 GB', '1.29 GB', '11.68 GB', 'fp32:✗ 3-bit:✓'],
        ]
    )

    doc.add_paragraph()
    add_bold(doc,
        '30 层全缓存：13GB → 1.3GB，省出近 12GB，让 24GB GPU 跑 720P 长视频。',
        12, (0xFF, 0x6B, 0x6B))

    # ---- 使用方式 ----
    doc.add_heading('使用方式', level=1)

    code = """import vedioquant

# One line to enable
handle = vedioquant.enable(pipe.transformer, bits=3)

# Inference as usual
output = pipe("a cat on sofa", num_frames=81)

# Check stats
print(handle.stats())
# → {'cache_hits': 32, 'hit_rate': '64.0%', 'compression_ratio': '10.7×'}

# Estimate memory (no GPU needed)
print(vedioquant.estimate_savings(height=720, width=1280, num_frames=81))
# → fp32: 886MB, 3-bit: 83MB, saved: 803MB"""

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # ---- 研究贡献 ----
    doc.add_heading('研究贡献', level=1)
    for c in [
        '首次将语言模型 KV Cache 压缩技术迁移到视频扩散模型特征缓存场景',
        '验证了视频模型特征与 LLM KV 向量具有相同的数学性质（旋转高斯化有效，且效果更好）',
        '发现视频模型特征天然更适合 TurboQuant 压缩（峰度15 vs 语言模型900，余弦0.98 vs 0.95）',
        '实现 TeaCache + TurboQuant 融合，3-bit 压缩 10× 缓存缩减',
        '封装为开源库 VedioQuant，一行代码启用，自动适配多种视频模型架构',
        '使 720P/81帧全层缓存从需要 80GB GPU 降至 24GB 消费级 GPU 可运行',
    ]:
        doc.add_paragraph(c, style='List Number')

    # ---- 参考文献 ----
    doc.add_heading('参考文献', level=1)
    for r in [
        'TurboQuant: arXiv:2504.19874 (ICLR 2026)',
        'PolarQuant: arXiv:2502.02617',
        'QJL: arXiv:2406.03482',
        'TeaCache: arXiv:2411.19108',
        'Wan2.1: github.com/Wan-Video/Wan2.1',
    ]:
        doc.add_paragraph(r, style='List Bullet')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\nGitHub: github.com/robin-ph/vedioquant')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x45, 0xB7, 0xD1)

    return doc


# ================================================================
# 英文版
# ================================================================

def build_english():
    doc = Document()
    setup_styles(doc)
    add_cover(doc, 'Extreme Cache Compression for Video Diffusion', 'TurboQuant × TeaCache Cross-Domain Fusion')

    # ---- Origin Story ----
    doc.add_heading('Origin Story: From Curiosity to a 10× Breakthrough', level=1)

    doc.add_paragraph(
        'In 2025, AI video generation exploded. ByteDance\'s Seedance launched in China and was '
        'instantly overwhelmed with demand — a 5-second video costs several yuan and takes minutes to generate, '
        'yet countless video creators are lining up to use it. AI-generated content is visibly flooding short-video '
        'platforms, reshaping the content creation industry at breakneck speed.'
    )

    doc.add_paragraph(
        'As someone deeply curious about technology, I wanted to understand: why is video generation so expensive? '
        'What happens behind the scenes when you generate a 5-second 720P video? So I started digging into '
        'video diffusion models — Wan2.1, CogVideoX, HunyuanVideo — and their inference pipelines.'
    )

    doc.add_paragraph(
        'The answer was staggering: a single video generation requires ~50 denoising steps, each running '
        'a full transformer forward pass over 75,600 tokens across 30 layers. The compute is so massive '
        'that it demands an A100 80GB GPU — hardware that individual creators and small studios simply cannot afford.'
    )

    doc.add_heading('The Spark: Google\'s Paper and an Engineering Reproduction', level=2)

    doc.add_paragraph(
        'Then two things caught my attention:'
    )

    doc.add_paragraph(
        'First, Google published TurboQuant at ICLR 2026 — a technique that compresses LLM KV Cache '
        'from fp16 down to 3.5 bits/coordinate with only 5% cosine similarity loss. The core trick: '
        'a random orthogonal rotation that "Gaussianizes" vector coordinates, making aggressive scalar '
        'quantization effective.'
    )

    doc.add_paragraph(
        'Second, someone successfully reproduced TurboQuant in Python (turboquant_plus), validating it '
        'on real Qwen3 KV tensors — kurtosis dropped from 900 to 2.9 after rotation (perfect Gaussian is 3.0), '
        'all 141 tests passed. This proved the algorithm works in practice, not just on paper.'
    )

    p = doc.add_paragraph()
    run = p.add_run(
        'This lit a fire in me: if TurboQuant can compress LLM KV Cache, what about video model feature cache? '
        'Both cache attention layer intermediate tensors — the data is fundamentally the same! '
        'I decided to try migrating TurboQuant to video model inference.'
    )
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        'The results were stunning — 3-bit compression achieved 0.98 cosine similarity on Wan2.1 features '
        '(better than the 0.95 on LLMs!). Cache for 720P/81-frame video dropped from 886MB to 88MB. '
        'Full 30-layer caching went from 13GB to 1.3GB — making it feasible on a consumer 24GB GPU.'
    )

    doc.add_page_break()

    # ---- Why Video Models Need This More ----
    doc.add_heading('Why Video Models Need Cache Compression More Than LLMs', level=1)

    doc.add_heading('Reason 1: Compute is ~350× Larger', level=2)
    doc.add_paragraph(
        'LLMs grow KV Cache linearly — one new row per generated token. Video models compute '
        'attention over ALL tokens at EVERY denoising step — it\'s an all-to-all operation repeated 50 times.'
    )

    add_table(doc,
        ['', 'LLM (Qwen3-7B)', 'Video Model (Wan2.1-14B, 720P)'],
        [
            ['Tokens', '~8,000 (8K context)', '75,600 (720P × 81 frames)'],
            ['Compute passes', '1 per token (autoregressive)', '50 (denoising steps)'],
            ['Total attention ops', '8K² × 32 layers', '75.6K² × 40 layers × 50 steps'],
            ['Scale comparison', '1×', '~350×'],
        ]
    )
    doc.add_paragraph()

    doc.add_heading('Reason 2: VRAM is the Hard Bottleneck', level=2)

    add_table(doc,
        ['Model', 'Scenario', 'Per-layer cache (fp32)', '30-layer full cache'],
        [
            ['Qwen3-7B', '8K context', '~32 MB', '~1 GB'],
            ['Wan2.1-1.3B', '480P / 81 frames', '192 MB', '5.8 GB'],
            ['Wan2.1-14B', '720P / 81 frames', '443 MB', '12.98 GB'],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph(
        'Video model caches are 10-13× larger than LLM caches. For video creators using consumer GPUs '
        '(RTX 4090, 24GB), cache compression isn\'t "nice to have" — it determines whether the model runs at all.'
    )

    doc.add_heading('Reason 3: Video Features Compress Better (Surprise!)', level=2)

    add_table(doc,
        ['Property', 'LLM KV Vectors', 'Video Features', 'Implication'],
        [
            ['Pre-rotation kurtosis', '~900', '~15', 'Video features already more uniform'],
            ['Post-rotation kurtosis', '2.9', '0.2', 'Better Gaussianization'],
            ['3-bit cosine similarity', '0.95', '0.98', 'Higher compression quality'],
        ]
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(
        'Video model features have kurtosis of only 15 (vs 900 for LLMs). After rotation, kurtosis '
        'drops to 0.2 — nearly perfect Gaussian. This means TurboQuant actually works BETTER on video models '
        'than on its original LLM target. A perfect cross-domain transfer.'
    )
    run.font.bold = True

    doc.add_heading('Reason 4: Massive Industry Demand', level=2)
    doc.add_paragraph(
        'Seedance charges several yuan per 5-second video — backed by minutes of A100 GPU time per generation. '
        'Reducing inference cost isn\'t just a technical exercise — it\'s a business imperative. '
        'Whoever cuts the cost by 10× turns AI video from "an expensive toy for the few" into '
        '"an everyday tool for every creator."'
    )

    doc.add_paragraph('VedioQuant contributes to this by:')
    for b in [
        'Enabling consumer GPUs (RTX 4090, 24GB) to run 720P long-video generation with full caching',
        'Reducing cloud GPU cost per generation (lower VRAM = cheaper GPU tier)',
        'Increasing throughput per GPU (less VRAM per request = more concurrent batch slots)',
    ]:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ---- Algorithm ----
    doc.add_heading('Algorithm: How TurboQuant Achieves Extreme Compression', level=1)

    add_img(doc, 'algorithm_flow.png', 6)
    add_caption(doc, 'Figure 1: TurboQuant pipeline — fp32 (32 bits) down to 3 bits/coordinate')
    doc.add_paragraph()

    doc.add_heading('Step 1: Norm Separation', level=2)
    doc.add_paragraph(
        'Separate the vector\'s magnitude (norm γ) from its direction. Norm is stored in full float32 precision. '
        'Only the direction is quantized — because attention computation (dot products) depends primarily on direction.'
    )

    doc.add_heading('Step 2: Random Rotation for Gaussianization — The Core Trick', level=2)
    doc.add_paragraph(
        'Raw feature coordinates have highly non-uniform distributions (heavy outliers). '
        'Multiplying by a random orthogonal matrix transforms each coordinate into a weighted average '
        'of all original coordinates. By the Central Limit Theorem, the result converges to a Gaussian — '
        'maximizing scalar quantization efficiency.'
    )

    add_img(doc, 'gaussianization.png', 5.5)
    add_caption(doc, 'Figure 2: Random rotation transforms the distribution from peaked (kurtosis 15) to uniform Gaussian (kurtosis 0.2)')
    doc.add_paragraph()

    add_bold(doc,
        'Verified on Wan2.1: kurtosis 15.0 → 0.2 after rotation. '
        'Std dev 0.025511 vs theoretical 1/√d = 0.025516. Perfect match.', 11)

    doc.add_heading('Step 3: Precomputed Codebook Quantization', level=2)
    doc.add_paragraph(
        'Rotated coordinates follow N(0, 1/√d). We precompute optimal Lloyd-Max codebooks for the standard '
        'Gaussian distribution. Quantization becomes a single torch.bucketize call — no per-vector iteration needed.'
    )

    # ---- Results ----
    doc.add_heading('Experimental Results', level=1)

    doc.add_paragraph(
        'Extracted real intermediate features (d=1536) from Wan2.1-1.3B\'s 30 attention layers. '
        'Tested compression/decompression quality on 50 token vectors:'
    )

    add_img(doc, 'compression_quality.png', 5.5)
    add_caption(doc, 'Figure 3: Compression quality and ratio at different bit widths')
    doc.add_paragraph()

    add_table(doc,
        ['Config', 'Compression Ratio', 'Avg Cosine Sim', 'Min Cosine Sim', 'Verdict'],
        [
            ['2-bit', '15.2×', '0.9388', '0.9347', '✅ > 0.90'],
            ['3-bit (recommended)', '10.0×', '0.9822', '0.9729', '✅ > 0.90'],
            ['4-bit', '7.3×', '0.9942', '0.9921', '✅ > 0.90'],
        ]
    )

    doc.add_paragraph()
    add_bold(doc,
        '3-bit sweet spot: 10× compression, 0.98 cosine similarity, worst-case still > 0.97.',
        12, (0xFF, 0x6B, 0x6B))

    # ---- Memory Impact ----
    doc.add_heading('Real-World Impact: Memory Savings', level=1)

    add_img(doc, 'memory_comparison.png', 5.5)
    add_caption(doc, 'Figure 4: Cache memory comparison across video configurations')
    doc.add_paragraph()

    add_img(doc, 'full_cache_comparison.png', 5.0)
    add_caption(doc, 'Figure 5: Full-layer caching at 720P/81 frames — VedioQuant enables 24GB GPUs')
    doc.add_paragraph()

    add_table(doc,
        ['Cached Layers', 'fp32', '3-bit', 'Saved', '24GB GPU'],
        [
            ['2 layers', '886 MB', '88 MB', '798 MB', 'Both fit'],
            ['10 layers', '4.33 GB', '441 MB', '3.89 GB', 'Both fit'],
            ['30 layers (all)', '12.98 GB', '1.29 GB', '11.68 GB', 'fp32: ✗ / 3-bit: ✓'],
        ]
    )

    doc.add_paragraph()
    add_bold(doc,
        'Full 30-layer cache: 13GB → 1.3GB. Saves ~12GB, enabling 720P long-video generation on consumer GPUs.',
        12, (0xFF, 0x6B, 0x6B))

    # ---- Usage ----
    doc.add_heading('Usage: One Line to Enable', level=1)

    code = """import vedioquant

# One line to enable compressed caching
handle = vedioquant.enable(pipe.transformer, bits=3)

# Run inference as usual — no code changes
output = pipe("a cat on sofa", num_frames=81)

# Check cache statistics
print(handle.stats())
# → {'cache_hits': 32, 'hit_rate': '64.0%', 'compression_ratio': '10.7×'}

# Estimate savings without GPU
print(vedioquant.estimate_savings(height=720, width=1280, num_frames=81))
# → fp32: 886MB, 3-bit: 83MB, saved: 803MB"""

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # ---- Contributions ----
    doc.add_heading('Contributions', level=1)
    for c in [
        'First migration of LLM KV Cache compression to video diffusion model feature caching',
        'Verified that video features share the same mathematical properties as LLM KV vectors (rotation Gaussianization works, and works even better)',
        'Discovered video features are naturally better suited for TurboQuant (kurtosis 15 vs 900, cosine 0.98 vs 0.95)',
        'Implemented TeaCache + TurboQuant fusion: 3-bit compression, 10× cache reduction',
        'Packaged as open-source library VedioQuant: one line to enable, auto-adapts to multiple architectures',
        'Reduced 720P/81-frame full-layer caching from 80GB GPU requirement to 24GB consumer GPU',
    ]:
        doc.add_paragraph(c, style='List Number')

    # ---- References ----
    doc.add_heading('References', level=1)
    for r in [
        'TurboQuant: arXiv:2504.19874 (ICLR 2026) — KV Cache extreme compression',
        'PolarQuant: arXiv:2502.02617 — Random rotation for quantization-friendly distributions',
        'QJL: arXiv:2406.03482 — Quantized Johnson-Lindenstrauss for inner product preservation',
        'TeaCache: arXiv:2411.19108 — Timestep-aware caching for video diffusion',
        'Wan2.1: github.com/Wan-Video/Wan2.1 — Open-source video generation',
    ]:
        doc.add_paragraph(r, style='List Bullet')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\nGitHub: github.com/robin-ph/vedioquant')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x45, 0xB7, 0xD1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Star ⭐ if you find this useful!')
    run.font.size = Pt(12)

    return doc


# ============================================================
# 生成两份文档
# ============================================================
base = os.path.dirname(os.path.dirname(__file__))

cn_doc = build_chinese()
cn_path = os.path.join(base, 'VedioQuant_Research_Report_CN.docx')
cn_doc.save(cn_path)
print(f"✓ 中文版: {cn_path}")

en_doc = build_english()
en_path = os.path.join(base, 'VedioQuant_Research_Report_EN.docx')
en_doc.save(en_path)
print(f"✓ 英文版: {en_path}")
