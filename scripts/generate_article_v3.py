"""
生成 VedioQuant 研究报告 v3
- 所有表格转为图片
- 生成 Twitter Banner
- 中文版 + 英文版
"""

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

assets_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets')
os.makedirs(assets_dir, exist_ok=True)

plt.rcParams.update({
    'font.size': 12,
    'figure.facecolor': 'white',
    'axes.facecolor': '#FAFAFA',
})


# ============================================================
# 表格渲染为图片
# ============================================================

def render_table_image(headers, rows, filename, title=None, highlight_row=None, col_widths=None):
    """把表格渲染成好看的图片"""
    n_rows = len(rows) + 1  # +1 header
    n_cols = len(headers)

    if col_widths is None:
        col_widths = [max(len(str(h)), max(len(str(r[i])) for r in rows)) for i, h in enumerate(headers)]
        total = sum(col_widths)
        col_widths = [w / total for w in col_widths]

    fig_width = max(8, n_cols * 2.2)
    fig_height = 0.5 + n_rows * 0.45 + (0.4 if title else 0)

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('off')
    ax.set_xlim(0, 1)
    ax.set_ylim(0, n_rows + (1 if title else 0))

    y_offset = 0
    if title:
        ax.text(0.5, n_rows + 0.3, title, ha='center', va='center',
                fontsize=14, fontweight='bold', color='#1A1A2E')
        y_offset = 0

    # 绘制表格
    for row_idx in range(n_rows):
        y = n_rows - row_idx - 1

        # 背景色
        if row_idx == 0:
            bg_color = '#2C3E50'
            text_color = 'white'
            fontweight = 'bold'
        elif highlight_row is not None and row_idx - 1 == highlight_row:
            bg_color = '#E8F8F5'
            text_color = '#1A1A2E'
            fontweight = 'bold'
        elif row_idx % 2 == 0:
            bg_color = '#F8F9FA'
            text_color = '#333333'
            fontweight = 'normal'
        else:
            bg_color = 'white'
            text_color = '#333333'
            fontweight = 'normal'

        # 画行背景
        rect = plt.Rectangle((0, y), 1, 0.9, facecolor=bg_color,
                              edgecolor='#DEE2E6', linewidth=0.5)
        ax.add_patch(rect)

        # 写文字
        x_pos = 0
        for col_idx in range(n_cols):
            w = col_widths[col_idx] if col_widths else 1.0 / n_cols
            cell_x = x_pos + w / 2

            if row_idx == 0:
                text = headers[col_idx]
            else:
                text = str(rows[row_idx - 1][col_idx])

            ax.text(cell_x, y + 0.45, text, ha='center', va='center',
                    fontsize=11, fontweight=fontweight, color=text_color)
            x_pos += w

    plt.tight_layout(pad=0.2)
    path = os.path.join(assets_dir, filename)
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f"  ✓ {filename}")
    return path


# ============================================================
# 生成所有表格图片
# ============================================================

print("生成表格图片...")

# 表1: 核心数据
render_table_image(
    ['Metric', 'Result'],
    [
        ['Cache Compression Ratio', '10.7× (fp32 → 3-bit)'],
        ['Cosine Similarity', '0.9822 (< 2% quality loss)'],
        ['720P 81-frame Cache', '886 MB → 83 MB'],
        ['Decompression Latency', '< 1ms per cache hit'],
    ],
    'table_key_results.png',
    title='VedioQuant Key Results',
    col_widths=[0.4, 0.6],
)

# 表2: 计算量对比
render_table_image(
    ['', 'LLM (Qwen3-7B)', 'Video (Wan2.1-1.3B, 720P)'],
    [
        ['Tokens', '~8,000', '75,600'],
        ['Compute Passes', '1 per token', '50 (denoising steps)'],
        ['Total Attention', '8K² x 32 layers', '75.6K² x 30L x 50 steps'],
        ['Scale', '1x', '~4000x'],
    ],
    'table_compute_comparison.png',
    title='Compute: LLM vs Video Model',
    col_widths=[0.25, 0.35, 0.40],
)

# 表3: 缓存大小对比
render_table_image(
    ['Model', 'Scenario', 'Per-Layer Cache', '30-Layer Full Cache'],
    [
        ['Qwen3-7B', '8K context', '~32 MB', '~1 GB'],
        ['Wan2.1-1.3B', '480P / 81 frames', '192 MB', '5.8 GB'],
        ['Wan2.1 (720P est.)', '720P / 81 frames', '443 MB', '12.98 GB'],
    ],
    'table_cache_size.png',
    title='Cache Size: LLM vs Video Model',
    highlight_row=2,
    col_widths=[0.22, 0.28, 0.25, 0.25],
)

# 表4: 特征性质对比
render_table_image(
    ['Property', 'LLM KV Vectors', 'Video Features', 'Implication'],
    [
        ['Pre-rotation Kurtosis', '~900 *', '~15', 'Video more uniform'],
        ['Post-rotation Kurtosis', '2.9 *', '0.2', 'Better Gaussianization'],
        ['3-bit Cosine Sim', '0.95 *', '0.98', 'Higher quality'],
    ],
    'table_feature_comparison.png',
    title='Feature Properties (* = from turboquant_plus repo, not our experiment)',
    highlight_row=2,
    col_widths=[0.28, 0.22, 0.22, 0.28],
)

# 表5: 压缩质量
render_table_image(
    ['Config', 'Compression', 'Avg Cosine', 'Min Cosine', 'Verdict'],
    [
        ['2-bit', '15.2x', '0.9388', '0.9347', 'PASS > 0.90'],
        ['3-bit (recommended)', '10.0x', '0.9822', '0.9729', 'PASS > 0.90'],
        ['4-bit', '7.3x', '0.9942', '0.9921', 'PASS > 0.90'],
    ],
    'table_compression_quality.png',
    title='Compression Quality on Wan2.1 Real Features',
    highlight_row=1,
    col_widths=[0.28, 0.18, 0.18, 0.18, 0.18],
)

# 表6: 显存节省
render_table_image(
    ['Cached Layers', 'fp32', '3-bit VedioQuant', 'Saved', '24GB GPU'],
    [
        ['2 layers', '886 MB', '83 MB', '803 MB', 'Both fit'],
        ['10 layers', '4.33 GB', '417 MB', '3.92 GB', 'Both fit'],
        ['30 layers (all)', '12.98 GB', '1.22 GB', '11.76 GB', 'fp32: NO / 3-bit: YES'],
    ],
    'table_memory_savings.png',
    title='Memory Savings at 720P / 81 Frames',
    highlight_row=2,
    col_widths=[0.20, 0.18, 0.22, 0.20, 0.20],
)

# 表7: 跨领域对比
render_table_image(
    ['', 'LLM KV Cache', 'Video Model TeaCache'],
    [
        ['Cached Content', 'Historical K/V vectors', 'Attention features across steps'],
        ['Purpose', 'Avoid recomputing history', 'Skip redundant denoising steps'],
        ['Data Type', 'Attention intermediate tensors', 'Attention intermediate tensors'],
        ['Compression Need', 'Long sequences → large cache', 'High-res video → large cache'],
    ],
    'table_cross_domain.png',
    title='Cross-Domain Analogy: Same Data, Same Solution',
    col_widths=[0.25, 0.35, 0.40],
)

# 表8: 支持模型
render_table_image(
    ['Model', 'Status', 'Detection Pattern'],
    [
        ['Wan2.1 (1.3B / 14B)', '[Verified]', 'model.blocks'],
        ['CogVideoX', '[Compatible]', 'model.transformer_blocks'],
        ['HunyuanVideo', '[Compatible]', 'model.transformer_blocks'],
        ['Any diffusers model', '[Auto-detect]', 'Multiple patterns'],
    ],
    'table_supported_models.png',
    title='Supported Models',
    col_widths=[0.35, 0.25, 0.40],
)


# ============================================================
# Twitter Banner (1500×500)
# ============================================================

print("\n生成 Twitter Banner...")

fig, ax = plt.subplots(figsize=(15, 5))
ax.set_xlim(0, 15)
ax.set_ylim(0, 5)
ax.axis('off')

# 背景渐变（用多个矩形模拟）
for i in range(150):
    x = i / 10
    r = int(26 + (44 - 26) * (i / 150))
    g = int(26 + (62 - 26) * (i / 150))
    b = int(46 + (80 - 46) * (i / 150))
    rect = plt.Rectangle((x, 0), 0.11, 5, facecolor=f'#{r:02x}{g:02x}{b:02x}', linewidth=0)
    ax.add_patch(rect)

# 装饰线条
for y in [0.8, 4.2]:
    ax.plot([1, 14], [y, y], color='#4ECDC4', linewidth=2, alpha=0.5)

# 标题
ax.text(7.5, 3.5, 'VedioQuant', ha='center', va='center',
        fontsize=52, fontweight='bold', color='white',
        fontfamily='Arial')

# 副标题
ax.text(7.5, 2.5, '10× Cache Compression for Video Diffusion Models',
        ha='center', va='center', fontsize=20, color='#4ECDC4',
        fontfamily='Arial')

# 数据亮点
highlights = [
    ('10.7×', 'Compression'),
    ('0.98', 'Cosine Sim'),
    ('886→88MB', 'Cache Size'),
    ('<1ms', 'Latency'),
]
for i, (val, label) in enumerate(highlights):
    x = 2.5 + i * 3.2
    # 数字
    ax.text(x, 1.5, val, ha='center', va='center',
            fontsize=24, fontweight='bold', color='#FF6B6B',
            fontfamily='Arial')
    # 标签
    ax.text(x, 1.0, label, ha='center', va='center',
            fontsize=12, color='#AAAAAA', fontfamily='Arial')

banner_path = os.path.join(assets_dir, 'twitter_banner.png')
plt.savefig(banner_path, dpi=200, bbox_inches='tight', facecolor='#1A1A2E')
plt.close()
print(f"  ✓ twitter_banner.png (3000×1000px)")


# ============================================================
# 文档生成函数
# ============================================================

def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3
    for i in range(1, 4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Arial'
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_img(doc, name, width=5.5):
    path = os.path.join(assets_dir, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def add_bold(doc, text, size=12, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_cover(doc, title_line1, title_line2, subtitle):
    for _ in range(1):
        doc.add_paragraph()
    # 插入 banner
    add_img(doc, 'twitter_banner.png', 6.5)
    doc.add_paragraph()

    # 文章标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_line1)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if title_line2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_line2)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x6B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


# ============================================================
# 中文版
# ============================================================

print("\n生成中文版文档...")

doc = Document()
setup_styles(doc)
add_cover(doc,
    '我把Google的LLM黑科技用在了阿里Wan2.1视频模型上',
    '缓存直降10倍，效果比语言模型上还好——AI视频的算力账单要变天了',
    'VedioQuant: TurboQuant × TeaCache 跨领域融合 | 在阿里Wan2.1上验证 | github.com/robin-ph/vedioquant'
)

doc.add_heading('起因：一个视频模型爱好者的好奇心', level=1)

doc.add_paragraph(
    '2025 年，AI 视频生成赛道彻底爆发。字节跳动的 Seedance 在中国一经发布就供不应求——'
    '一条 5 秒视频生成要等好几分钟，价格高昂，但仍有无数视频工作者和创作者抢着用。'
    '短视频平台上用 AI 生成的内容肉眼可见地增多，这个领域正在以惊人的速度重塑内容创作产业。'
)

doc.add_paragraph(
    '作为一个视频模型爱好者，我非常好奇：为什么视频生成这么贵？'
    '一段 5 秒 720P 的视频，背后到底发生了多少计算？'
    '于是我开始深入研究视频扩散模型的推理流程，特别是阿里开源的 Wan2.1。'
)

doc.add_paragraph(
    '答案让我震惊：一次视频生成需要 50 步去噪迭代，每步对 75,600 个 token 做 30 层 transformer 前向传播。'
    '计算量大到需要 A100 80GB 这样的顶级 GPU 才能跑起来。'
)

doc.add_heading('转折：Google 的论文和一个惊喜发现', level=2)

doc.add_paragraph(
    '就在这时，我看到 Google Research 发布了 TurboQuant（arXiv:2504.19874）——把大语言模型 KV Cache '
    '从 fp16 压缩到 3.5-bit，4× 压缩下余弦相似度仍有 0.95。核心 trick 是随机正交旋转让向量坐标高斯化。'
)

doc.add_paragraph(
    '紧接着，有人用 Python 工程复现了 TurboQuant（turboquant_plus，GitHub 1.7k stars），'
    '在 Qwen3 上验证了旋转后峰度从 900 降到 2.9。这证明了算法工程上完全可行。'
)

add_bold(doc,
    '这极大地刺激了我：语言模型和视频模型缓存的都是注意力层中间张量——数据本质完全相同！'
    '如果 TurboQuant 能压缩 KV Cache，那视频模型的特征缓存呢？', 12)

doc.add_paragraph(
    '于是我动手尝试。没想到效果这么惊人——'
)

add_img(doc, 'table_key_results.png', 5)
doc.add_paragraph()

doc.add_page_break()

# ---- 为什么视频模型更需要 ----
doc.add_heading('为什么视频模型比语言模型更需要缓存压缩？', level=1)

doc.add_heading('原因 1：计算量是语言模型的数千倍', level=2)
doc.add_paragraph(
    '语言模型 KV Cache 线性增长，视频模型每步对所有 token 做全量注意力，重复 50 次。'
    '粗略估算，视频模型的注意力计算量约为语言模型的 ~4000 倍。'
)
add_img(doc, 'table_compute_comparison.png', 5.5)
doc.add_paragraph()

doc.add_heading('原因 2：显存是最硬的瓶颈', level=2)
doc.add_paragraph(
    '视频模型的缓存比语言模型大 10-13×，直接决定能不能在消费级 GPU 上跑。'
)
add_img(doc, 'table_cache_size.png', 5.5)
doc.add_paragraph()

doc.add_heading('原因 3：视频特征天然更适合压缩', level=2)
doc.add_paragraph(
    '实验中的意外惊喜——视频特征的峰度只有 15（据 turboquant_plus 报告，语言模型约 900），'
    '旋转后几乎完美高斯（峰度 0.2 vs 理论 0.0），压缩质量比语言模型上还好。'
)
add_img(doc, 'table_feature_comparison.png', 5.5)
doc.add_paragraph()

doc.add_heading('原因 4：产业需求极其迫切', level=2)
doc.add_paragraph(
    'Seedance 一条 5 秒视频售价数元，背后是 A100 GPU 数分钟的成本。'
    '降低推理成本不仅是技术问题，更是商业问题——'
    '谁能把成本降一个数量级，谁就能让 AI 视频从"少数人的昂贵玩具"变成"每个创作者的日常工具"。'
)

doc.add_page_break()

# ---- 跨领域迁移 ----
doc.add_heading('跨领域迁移：为什么可行？', level=1)
doc.add_paragraph('语言模型 KV Cache 和视频模型 TeaCache 的本质对比：')
add_img(doc, 'table_cross_domain.png', 5.5)
add_bold(doc, '两者缓存的都是注意力层中间张量——TurboQuant 可以直接迁移。', 12, (0x4E, 0xCD, 0xC4))

doc.add_paragraph()

# ---- 算法 ----
doc.add_heading('算法原理', level=1)
add_img(doc, 'algorithm_flow.png', 6)
add_caption(doc, '图1: TurboQuant 压缩流程')
doc.add_paragraph()

doc.add_heading('Step 1: 范数分离', level=2)
doc.add_paragraph('长度和方向分开。范数 float32 全精度存，只对方向量化。')

doc.add_heading('Step 2: 随机旋转高斯化（核心 Trick）', level=2)
doc.add_paragraph(
    '原始坐标分布不均匀，随机正交旋转后趋近高斯分布（中心极限定理），标量量化效率最大化。'
)
add_img(doc, 'gaussianization.png', 5.5)
add_caption(doc, '图2: 旋转前后分布对比')
add_bold(doc, 'Wan2.1 实测：峰度 15.0 → 0.2，标准差完全吻合理论值。', 11)

doc.add_heading('Step 3: 预计算码本量化', level=2)
doc.add_paragraph('旋转后服从 N(0, 1/√d)，预计算码本 + torch.bucketize 查表，零迭代开销。')

doc.add_page_break()

# ---- 实验 ----
doc.add_heading('实验验证', level=1)
doc.add_paragraph('从 Wan2.1-1.3B 的 30 层注意力层提取真实特征（d=1536），50 个 token 测试：')
add_img(doc, 'compression_quality.png', 5.5)
add_caption(doc, '图3: 压缩质量和压缩比')
doc.add_paragraph()
add_img(doc, 'table_compression_quality.png', 5.5)
add_bold(doc, '3-bit 甜点：10× 压缩，余弦 0.98，最差向量也 > 0.97。', 12, (0xFF, 0x6B, 0x6B))

# ---- 显存 ----
doc.add_heading('显存节省', level=1)
add_img(doc, 'memory_comparison.png', 5.5)
add_caption(doc, '图4: 不同视频配置下缓存大小对比')
doc.add_paragraph()
add_img(doc, 'full_cache_comparison.png', 5.0)
add_caption(doc, '图5: 720P/81帧全层缓存')
doc.add_paragraph()
add_img(doc, 'table_memory_savings.png', 5.5)
add_bold(doc, '30 层全缓存：13GB → 1.2GB，省出近 12GB，24GB GPU 跑 720P 长视频。', 12, (0xFF, 0x6B, 0x6B))

doc.add_page_break()

# ---- 使用 ----
doc.add_heading('使用方式：一行代码', level=1)
code = """import vedioquant

handle = vedioquant.enable(pipe.transformer, bits=3)
output = pipe("a cat on sofa", num_frames=81)
print(handle.stats())
# → {'cache_hits': 32, 'hit_rate': '64.0%', 'compression_ratio': '10.7×'}

vedioquant.estimate_savings(height=720, width=1280, num_frames=81)
# → fp32: 886 MB, 3-bit: 83 MB, saved: 803 MB"""
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph()
add_img(doc, 'table_supported_models.png', 5)

# ---- 贡献 ----
doc.add_heading('研究贡献', level=1)
for c in [
    '首次将 LLM KV Cache 压缩技术迁移到视频扩散模型特征缓存',
    '验证视频特征天然更适合 TurboQuant（峰度 15 vs 900，余弦 0.98 vs 0.95）',
    '实现 TeaCache + TurboQuant 融合，3-bit 压缩 10× 缓存缩减',
    '封装为开源库 VedioQuant，一行启用，自动适配多种架构',
    '720P/81帧全层缓存：80GB GPU → 24GB 消费级 GPU',
]:
    doc.add_paragraph(c, style='List Number')

doc.add_heading('参考文献', level=1)
for r in [
    'TurboQuant: arXiv:2504.19874 (Google Research, 2025)',
    'PolarQuant: arXiv:2502.02617',
    'QJL: arXiv:2406.03482',
    'TeaCache: arXiv:2411.19108',
    'Wan2.1: github.com/Wan-Video/Wan2.1',
]:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nGitHub: github.com/robin-ph/vedioquant')
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x45, 0xB7, 0xD1)

cn_path = os.path.join(os.path.dirname(assets_dir), 'VedioQuant_Research_Report_CN.docx')
doc.save(cn_path)
print(f"✓ 中文版: {cn_path}")


# ============================================================
# 英文版
# ============================================================

print("\n生成英文版文档...")

doc = Document()
setup_styles(doc)
add_cover(doc,
    'I Took Google\'s LLM Trick and Tested It on Alibaba\'s Wan2.1 Video Model',
    '10x Cache Reduction — And It Works Even Better Than on LLMs',
    'VedioQuant: TurboQuant x TeaCache Cross-Domain Fusion | Verified on Wan2.1 | github.com/robin-ph/vedioquant'
)

doc.add_heading('Origin Story: A Video Model Enthusiast\'s 10x Discovery', level=1)

doc.add_paragraph(
    'In 2025, AI video generation exploded. ByteDance\'s Seedance launched in China and was '
    'instantly overwhelmed — a 5-second video costs several yuan and takes minutes to generate, '
    'yet countless creators are lining up to use it. AI-generated content is visibly flooding '
    'short-video platforms, reshaping the content creation industry at breakneck speed.'
)

doc.add_paragraph(
    'As a video model enthusiast, I was deeply curious: why is video generation so expensive? '
    'I started studying Alibaba\'s open-source Wan2.1 model and found that a single generation '
    'requires ~50 denoising steps, each running a full transformer forward pass over 75,600 tokens '
    'across 30 layers. It demands an A100 80GB GPU — hardware most people can\'t afford.'
)

doc.add_heading('The Spark', level=2)

doc.add_paragraph(
    'Then I saw Google Research\'s TurboQuant (arXiv:2504.19874) — compressing LLM KV Cache from fp16 to 3.5 bits '
    'with 0.95 cosine similarity. The core trick: random orthogonal rotation to Gaussianize coordinates.'
)

doc.add_paragraph(
    'Someone reproduced it in Python (turboquant_plus, 1.7k GitHub stars) and validated on Qwen3: '
    'kurtosis dropped from 900 to 2.9 after rotation. It works in practice, not just on paper.'
)

add_bold(doc,
    'This sparked an idea: both LLMs and video models cache attention layer intermediate tensors — '
    'the data is fundamentally the same! What if TurboQuant works on video model caches too?', 12)

doc.add_paragraph('The results were stunning:')
add_img(doc, 'table_key_results.png', 5)

doc.add_page_break()

doc.add_heading('Why Video Models Need Cache Compression More Than LLMs', level=1)

doc.add_heading('Reason 1: ~4000x More Compute', level=2)
add_img(doc, 'table_compute_comparison.png', 5.5)
doc.add_paragraph()

doc.add_heading('Reason 2: VRAM is the Hard Bottleneck', level=2)
doc.add_paragraph('Video model caches are 10-13× larger than LLM caches.')
add_img(doc, 'table_cache_size.png', 5.5)
doc.add_paragraph()

doc.add_heading('Reason 3: Video Features Compress Better (Surprise!)', level=2)
doc.add_paragraph(
    'An unexpected discovery — video features have kurtosis of only 15 (vs ~900 for LLMs, '
    'as reported by the turboquant_plus repo). After rotation, kurtosis drops to 0.2 — nearly '
    'perfect Gaussian. TurboQuant actually works BETTER on video models than on its original LLM target.'
)
add_img(doc, 'table_feature_comparison.png', 5.5)
doc.add_paragraph()

doc.add_heading('Reason 4: Massive Industry Demand', level=2)
doc.add_paragraph(
    'Seedance charges several yuan per 5-second video — backed by minutes of A100 time. '
    'Whoever cuts inference cost by 10× turns AI video from "expensive toy" into "everyday tool."'
)

doc.add_page_break()

doc.add_heading('Cross-Domain Transfer: Why It Works', level=1)
add_img(doc, 'table_cross_domain.png', 5.5)
add_bold(doc, 'Both cache attention intermediate tensors — TurboQuant transfers directly.', 12, (0x4E, 0xCD, 0xC4))

doc.add_paragraph()
doc.add_heading('Algorithm', level=1)
add_img(doc, 'algorithm_flow.png', 6)
add_caption(doc, 'Figure 1: TurboQuant compression pipeline')
doc.add_paragraph()

doc.add_heading('Step 1: Norm Separation', level=2)
doc.add_paragraph('Separate magnitude (full precision) from direction (quantized). Attention depends on direction.')

doc.add_heading('Step 2: Random Rotation — The Core Trick', level=2)
doc.add_paragraph(
    'Raw coordinates are non-uniform. Random orthogonal rotation → each coordinate becomes '
    'a weighted average of all originals → Central Limit Theorem → Gaussian distribution → '
    'optimal scalar quantization.'
)
add_img(doc, 'gaussianization.png', 5.5)
add_caption(doc, 'Figure 2: Distribution before and after rotation')
add_bold(doc, 'Wan2.1: kurtosis 15.0 → 0.2. Std dev matches theory exactly.', 11)

doc.add_heading('Step 3: Precomputed Codebook', level=2)
doc.add_paragraph('Rotated coords follow N(0, 1/√d). Precompute Lloyd-Max codebook → torch.bucketize. Zero iteration.')

doc.add_page_break()

doc.add_heading('Experimental Results', level=1)
doc.add_paragraph('Real features from Wan2.1-1.3B, 30 attention layers, d=1536, 50 tokens:')
add_img(doc, 'compression_quality.png', 5.5)
add_caption(doc, 'Figure 3: Compression quality across bit widths')
doc.add_paragraph()
add_img(doc, 'table_compression_quality.png', 5.5)
add_bold(doc, '3-bit sweet spot: 10× compression, 0.98 cosine, worst case > 0.97.', 12, (0xFF, 0x6B, 0x6B))

doc.add_heading('Memory Savings', level=1)
add_img(doc, 'memory_comparison.png', 5.5)
add_caption(doc, 'Figure 4: Cache memory across video configurations')
doc.add_paragraph()
add_img(doc, 'full_cache_comparison.png', 5.0)
add_caption(doc, 'Figure 5: Full-layer caching — VedioQuant enables 24GB GPUs')
doc.add_paragraph()
add_img(doc, 'table_memory_savings.png', 5.5)
add_bold(doc, 'Full 30-layer cache: 13GB → 1.2GB. Saves ~12GB. Consumer GPU runs 720P long video.', 12, (0xFF, 0x6B, 0x6B))

doc.add_page_break()

doc.add_heading('Usage: One Line to Enable', level=1)
code = """import vedioquant

handle = vedioquant.enable(pipe.transformer, bits=3)
output = pipe("a cat on sofa", num_frames=81)
print(handle.stats())
# → {'cache_hits': 32, 'hit_rate': '64.0%', 'compression_ratio': '10.7×'}

vedioquant.estimate_savings(height=720, width=1280, num_frames=81)
# → fp32: 886 MB, 3-bit: 83 MB, saved: 803 MB"""
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)
doc.add_paragraph()
add_img(doc, 'table_supported_models.png', 5)

doc.add_heading('Contributions', level=1)
for c in [
    'First migration of LLM KV Cache compression to video diffusion feature caching',
    'Discovered video features naturally compress better (kurtosis 15 vs 900, cosine 0.98 vs 0.95)',
    'Implemented TeaCache + TurboQuant fusion: 3-bit, 10× cache reduction',
    'Open-source library VedioQuant: one line, auto-adapts to multiple architectures',
    '720P/81-frame full cache: 80GB GPU → 24GB consumer GPU',
]:
    doc.add_paragraph(c, style='List Number')

doc.add_heading('References', level=1)
for r in [
    'TurboQuant: arXiv:2504.19874 (Google Research, 2025)',
    'PolarQuant: arXiv:2502.02617',
    'QJL: arXiv:2406.03482',
    'TeaCache: arXiv:2411.19108',
    'Wan2.1: github.com/Wan-Video/Wan2.1',
]:
    doc.add_paragraph(r, style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nGitHub: github.com/robin-ph/vedioquant')
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x45, 0xB7, 0xD1)

en_path = os.path.join(os.path.dirname(assets_dir), 'VedioQuant_Research_Report_EN.docx')
doc.save(en_path)
print(f"✓ 英文版: {en_path}")

print(f"\n✓ Twitter Banner: {banner_path}")
print("\n全部完成!")
